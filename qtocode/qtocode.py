import re
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

# Function to process the Word document and generate Python code
def process_docx_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])

    if file_path:
        try:
            questions_and_answers = extract_questions_answers(file_path)
            code = generate_python_code(questions_and_answers)
            save_to_word_document(code)

            # Display a success message
            messagebox.showinfo("Success", "Document processed and saved successfully!")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")

# Function to extract questions and answers from the Word document
def extract_questions_answers(file_path):
    doc = Document(file_path)
    questions_and_answers = []

    current_question = None
    options = []
    correct_answer = None
    explanation = None
    in_question = False  # Track if we're currently parsing a question

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            # Check for a new question format (e.g., "1. **What does...")
            if re.match(r"^\d+\.", text):
                if current_question:
                    qa = {
                        "text": current_question,
                        "options": options,
                        "correctAnswer": correct_answer,
                        "explanation": explanation,
                    }
                    questions_and_answers.append(qa)

                current_question = text
                options = []
                correct_answer = None
                explanation = None
                in_question = True
            elif in_question and text.startswith("**"):
                # Append to the current question
                current_question += " " + text
            elif text.startswith(("a)", "b)", "c)", "d)")):
                # Remove the leading 'a)', 'b)', 'c)', 'd)' and strip white spaces
                option = re.sub(r"^[abcd]\)", "", text).strip()
                options.append(option)
                in_question = False  # No longer in question section
            elif text.startswith("**Correct Answer:"):
                # Extract and clean the correct answer
                correct_answer = re.sub(r"^[abcd]\)", "", text.replace("**Correct Answer:", "").replace("**", "").strip())
            elif text.startswith("Explanation:"):
                explanation = text.replace("Explanation:", "").strip()

    # Add the last question to the list
    if current_question:
        qa = {
            "text": current_question,
            "options": options,
            "correctAnswer": correct_answer,
            "explanation": explanation,
        }
        questions_and_answers.append(qa)

    return questions_and_answers

# Function to generate Python code from questions and answers
def generate_python_code(questions_and_answers):
    code = ""
    for qa in questions_and_answers:
        # Format the question text and options
        question_text = qa['text']
        options = ', '.join([f'"{o}"' for o in qa['options']])

        # Remove leading/trailing spaces from correct answer
        correct_answer = qa['correctAnswer'].strip()

        # Format the explanation
        explanation = qa['explanation']

        # Construct the Python code for this question
        code += f'Question(text: "{question_text}", options: [{options}], correctAnswer: "{correct_answer}", explanation: "{explanation}"),\n\n'
    return code



# Function to save generated Python code to a Word document
def save_to_word_document(code):
    output_doc = Document()
    output_doc.add_heading("Generated Python Code", level=1)
    output_doc.add_paragraph(code)

    output_file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])

    if output_file_path:
        output_doc.save(output_file_path)

# Create the main application window
root = tk.Tk()
root.title("Word Document Processor")

# Create a button to open a file dialog
open_file_button = tk.Button(root, text="Open Word Document", command=process_docx_file)
open_file_button.pack(pady=10)

# Start the GUI application
root.mainloop() 